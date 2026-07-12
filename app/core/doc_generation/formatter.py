"""
文档格式化工具

提供 Word 文档的字体、段落格式设置功能
"""
import logging
from typing import Optional
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from .config import (
    FONT_CHINESE_BODY,
    FONT_CHINESE_HEADING,
    FONT_ENGLISH,
    FONT_SIZE_BODY,
    FONT_SIZE_HEADING_1,
    FONT_SIZE_HEADING_2,
    FONT_SIZE_HEADING_3,
    LINE_SPACING,
    SPACE_BEFORE,
    FIRST_LINE_INDENT,
    COLOR_BLACK,
)

logger = logging.getLogger(__name__)


def set_font(run, font_name: str, size: Optional[Pt] = None, bold: bool = False):
    """
    设置 run 的字体（支持中英文混排）
    
    Args:
        run: docx 的 run 对象
        font_name: 字体名称
        size: 字体大小
        bold: 是否加粗
    """
    # 确保 rPr 元素存在
    rPr = run._element.get_or_add_rPr()
    
    # 确保 rFonts 元素存在
    rFonts = rPr.get_or_add_rFonts()
    
    # 设置字体
    rFonts.set(qn('w:eastAsia'), font_name)  # 中文字体
    rFonts.set(qn('w:ascii'), font_name)     # 英文字体
    rFonts.set(qn('w:hAnsi'), font_name)     # 复杂脚本字体
    
    # 设置颜色为黑色
    run.font.color.rgb = COLOR_BLACK
    
    if size:
        run.font.size = size
    if bold:
        run.font.bold = True


def set_paragraph_font(paragraph, font_name: str, size: Optional[Pt] = None, bold: bool = False):
    """
    设置段落中所有 run 的字体
    
    Args:
        paragraph: docx 的 paragraph 对象
        font_name: 字体名称
        size: 字体大小
        bold: 是否加粗
    """
    # 如果段落没有 runs，创建一个
    if not paragraph.runs:
        paragraph.add_run('')
    
    for run in paragraph.runs:
        set_font(run, font_name, size, bold)


def set_paragraph_format(paragraph, line_spacing: float = 1.5, space_before: Optional[Pt] = None, 
                        first_line_indent: bool = True, justify: bool = False):
    """
    设置段落格式（行距、段前段后距、首行缩进、对齐方式）
    
    Args:
        paragraph: docx 的 paragraph 对象
        line_spacing: 行距倍数（默认 1.5）
        space_before: 段前距（Pt 单位）
        first_line_indent: 是否首行缩进 2 字符（默认 True）
        justify: 是否两端对齐（默认 False，避免覆盖标题对齐设置）
    """
    # 设置行距（1.5 倍）
    paragraph.paragraph_format.line_spacing_rule = 1  # 1 = 倍数行距
    paragraph.paragraph_format.line_spacing = Pt(line_spacing * 12)  # 基于 12pt 字号
    
    # 设置段前距
    if space_before:
        paragraph.paragraph_format.space_before = space_before
    
    # 设置首行缩进（2 字符）
    if first_line_indent:
        paragraph.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    
    # 设置两端对齐（仅当 justify=True 时）
    if justify:
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def setup_document_style(doc: Document) -> None:
    """
    设置文档默认样式
    
    Args:
        doc: Document 对象
    """
    style = doc.styles['Normal']
    
    # 设置字体
    style.font.name = FONT_ENGLISH
    style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CHINESE_BODY)
    style._element.rPr.rFonts.set(qn('w:ascii'), FONT_ENGLISH)
    style._element.rPr.rFonts.set(qn('w:hAnsi'), FONT_ENGLISH)
    style.font.size = FONT_SIZE_BODY  # 小四号 (12pt)
    
    # 设置颜色为黑色
    style.font.color.rgb = COLOR_BLACK
    
    # 设置段落格式
    style.paragraph_format.line_spacing_rule = 1  # 倍数行距
    style.paragraph_format.line_spacing = Pt(LINE_SPACING * 12)
    style.paragraph_format.space_before = SPACE_BEFORE  # 段前 0.5 行
    style.paragraph_format.first_line_indent = FIRST_LINE_INDENT  # 首行缩进 2 字符
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 两端对齐


def add_formatted_heading(doc: Document, text: str, level: int = 1):
    """
    添加格式化的标题
    
    Args:
        doc: Document 对象
        text: 标题文本
        level: 标题级别 (1, 2, 3)
    
    Returns:
        标题段落对象
    """
    heading = doc.add_heading(text, level)
    
    # 一级标题居中，其他标题左对齐
    if level <= 1:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 根据级别设置字号
    size_map = {
        0: FONT_SIZE_HEADING_1,  # 主标题
        1: FONT_SIZE_HEADING_1,
        2: FONT_SIZE_HEADING_2,
        3: FONT_SIZE_HEADING_3,
    }
    size = size_map.get(level, FONT_SIZE_HEADING_3)
    
    set_paragraph_font(heading, FONT_CHINESE_HEADING, size, bold=True)
    # 标题不设置首行缩进
    set_paragraph_format(heading, LINE_SPACING, None, first_line_indent=False)
    return heading


def add_formatted_paragraph(doc: Document, text: str, style: str = None, 
                           font_name: str = None, size: Optional[Pt] = None, 
                           bold: bool = False, first_line_indent: bool = True) -> None:
    """
    添加格式化的段落
    
    Args:
        doc: Document 对象
        text: 段落文本
        style: 段落样式
        font_name: 字体名称
        size: 字体大小
        bold: 是否加粗
        first_line_indent: 是否首行缩进（默认 True）
    """
    if style:
        paragraph = doc.add_paragraph(text, style=style)
    else:
        paragraph = doc.add_paragraph(text)
    
    if font_name is None:
        font_name = FONT_CHINESE_BODY
    if size is None:
        size = FONT_SIZE_BODY
    
    set_paragraph_font(paragraph, font_name, size, bold)
    # 正文段落启用两端对齐
    set_paragraph_format(paragraph, LINE_SPACING, SPACE_BEFORE, first_line_indent, justify=True)
