"""
图片处理工具

提供 Markdown 图片解析和本地图片插入功能
"""
import os
import re
import tempfile
import logging
from typing import Optional, Tuple

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .config import IMAGE_SEARCH_PATHS, FONT_SIZE_CAPTION, LINE_SPACING, SPACE_BEFORE
from .formatter import set_paragraph_font, set_paragraph_format

logger = logging.getLogger(__name__)


def download_image(url: str) -> Optional[bytes]:
    """
    读取本地图片内容
    
    仅支持本地文件路径，不支持网络图片
    
    Args:
        url: 图片路径（可以是相对路径或绝对路径）
    
    Returns:
        图片二进制数据，失败返回 None
    """
    try:
        # 处理 file:// 前缀
        if url.startswith('file://'):
            local_path = url[7:]
        else:
            local_path = url
        
        # 如果是相对路径，尝试在多个目录查找
        if not os.path.isabs(local_path):
            # 尝试当前工作目录
            if os.path.exists(local_path):
                with open(local_path, 'rb') as f:
                    return f.read()
            
            # 尝试配置的搜索路径
            for base_dir in IMAGE_SEARCH_PATHS:
                full_path = os.path.join(base_dir, local_path)
                if os.path.exists(full_path):
                    with open(full_path, 'rb') as f:
                        return f.read()
            
            logger.warning(f"本地图片不存在：{local_path}")
            return None
        
        # 绝对路径
        local_path = os.path.expanduser(local_path)
        if os.path.exists(local_path):
            with open(local_path, 'rb') as f:
                return f.read()
        else:
            logger.warning(f"本地图片不存在：{local_path}")
            return None
        
    except Exception as e:
        logger.warning(f"读取图片失败 {url}: {e}")
        return None


def parse_markdown_image(line: str) -> Tuple[Optional[str], Optional[str], str]:
    """
    解析 Markdown 图片语法
    
    Args:
        line: 包含图片的文本行
    
    Returns:
        (alt_text, image_url, remaining_text)
        如果没有图片，返回 (None, None, original_line)
    """
    # 匹配 ![alt](url) 格式
    pattern = r'!\[([^\]]*)\]\(([^)]+)\)'
    match = re.search(pattern, line)
    
    if match:
        alt_text = match.group(1)
        image_url = match.group(2)
        remaining_text = line[:match.start()] + line[match.end():]
        return alt_text, image_url, remaining_text.strip()
    
    return None, None, line


def insert_image_to_doc(doc: Document, image_url: str, alt_text: str, 
                       image_counter: int) -> bool:
    """
    插入图片到 Word 文档
    
    Args:
        doc: Document 对象
        image_url: 图片路径
        alt_text: 图片说明文字
        image_counter: 图片计数器
    
    Returns:
        是否成功插入
    """
    image_data = download_image(image_url)
    
    if not image_data:
        # 下载失败，添加文字占位
        p = doc.add_paragraph(f"[图片下载失败：{alt_text}]")
        set_paragraph_font(p, FONT_CHINESE_BODY, FONT_SIZE_BODY)
        set_paragraph_format(p, LINE_SPACING, SPACE_BEFORE)
        return False
    
    try:
        # 将图片保存到临时文件
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
            tmp.write(image_data)
            tmp_path = tmp.name
        
        # 添加图片到文档
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(tmp_path, width=Cm(15))  # 限制图片宽度
        
        # 添加图片说明（如果有 alt 文本）
        if alt_text:
            caption = doc.add_paragraph(f"图{image_counter}: {alt_text}")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_font(caption, FONT_CHINESE_BODY, FONT_SIZE_CAPTION)
            set_paragraph_format(caption, LINE_SPACING, SPACE_BEFORE)
        
        # 清理临时文件
        try:
            os.unlink(tmp_path)
        except:
            pass
            
        logger.info(f"成功插入图片 {image_counter}: {image_url}")
        return True
        
    except Exception as e:
        logger.error(f"插入图片失败 {image_url}: {e}")
        # 插入失败时，添加文字说明
        p = doc.add_paragraph(f"[图片：{alt_text}]")
        set_paragraph_font(p, FONT_CHINESE_BODY, FONT_SIZE_BODY)
        set_paragraph_format(p, LINE_SPACING, SPACE_BEFORE)
        return False
