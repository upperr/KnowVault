"""
Markdown 到 Word 转换器

将 Markdown 格式内容转换为 Word 文档
"""
import io
import logging
from docx import Document

from .config import FONT_CHINESE_BODY, FONT_SIZE_BODY
from .formatter import (
    setup_document_style,
    add_formatted_heading,
    add_formatted_paragraph,
    set_paragraph_font,
    set_paragraph_format,
)
from .image_handler import parse_markdown_image, insert_image_to_doc
from .markdown_parser import parse_heading, parse_list_item, parse_inline_formatting

logger = logging.getLogger(__name__)


def markdown_to_docx(markdown_content: str, title: str = "文档") -> bytes:
    """
    将 Markdown 内容转换为 Word 文档
    
    功能：
    1. 解析 Markdown 标题、列表、段落
    2. 插入本地图片（支持相对路径和绝对路径）
    3. 设置字体：中文标题用黑体，正文用宋体，英文用 Times New Roman
    4. 设置格式：小四号字、1.5 倍行距、段前 0.5 行
    
    Args:
        markdown_content: Markdown 格式内容
        title: 文档标题
    
    Returns:
        Word 文档的二进制数据
    """
    doc = Document()
    
    # 设置文档默认样式
    setup_document_style(doc)
    
    # 添加标题
    add_formatted_heading(doc, title, level=0)
    doc.add_paragraph()  # 空行
    
    # 解析 Markdown
    lines = markdown_content.split('\n')
    current_list = []
    in_list = False
    image_counter = 0
    
    for line in lines:
        stripped = line.strip()
        
        # 空行处理
        if not stripped:
            _flush_list(doc, current_list)
            current_list = []
            in_list = False
            doc.add_paragraph()
            continue
        
        # 检查是否包含图片
        alt_text, image_url, remaining_text = parse_markdown_image(stripped)
        
        # 如果整行只是图片
        if alt_text is not None and not remaining_text.strip():
            if image_url:
                image_counter += 1
                insert_image_to_doc(doc, image_url, alt_text, image_counter)
            continue
        
        # 处理标题
        heading_level, heading_text = parse_heading(stripped)
        if heading_level:
            _flush_list(doc, current_list)
            current_list = []
            in_list = False
            
            # 如果标题中包含图片
            if alt_text and image_url:
                heading_text = remaining_text
            
            add_formatted_heading(doc, heading_text, level=heading_level)
            continue
        
        # 处理列表项
        is_list, list_text, list_type = parse_list_item(stripped)
        if is_list:
            in_list = True
            if alt_text and image_url:
                list_text = remaining_text
            current_list.append(parse_inline_formatting(list_text))
            continue
        
        # 结束列表，添加普通段落
        _flush_list(doc, current_list)
        current_list = []
        in_list = False
        
        paragraph_text = parse_inline_formatting(stripped)
        if alt_text and image_url:
            paragraph_text = parse_inline_formatting(remaining_text)
        
        add_formatted_paragraph(doc, paragraph_text)
    
    # 处理末尾的列表
    _flush_list(doc, current_list)
    
    # 保存文档
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _flush_list(doc: Document, items: list) -> None:
    """将列表项添加到文档"""
    if not items:
        return
    
    for item_text in items:
        p = doc.add_paragraph(item_text, style='List Bullet')
        set_paragraph_font(p, FONT_CHINESE_BODY, FONT_SIZE_BODY)
        # 列表项不设置段前距和首行缩进
        set_paragraph_format(p, 1.5, None, first_line_indent=False)
