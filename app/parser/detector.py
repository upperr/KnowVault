"""
文档检测模块
检测文件大小和页数，判断是否需要分割
"""
import os
import logging
from pathlib import Path
from typing import Optional

from app.parser.config import MAX_FILE_SIZE_MB, MAX_PAGES

logger = logging.getLogger(__name__)


def get_file_size_mb(file_path: str) -> float:
    """获取文件大小（MB）"""
    size_bytes = os.path.getsize(file_path)
    return size_bytes / (1024 * 1024)


def get_pdf_page_count(file_path: str) -> int:
    """获取 PDF 文件页数"""
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            return len(pdf.pages)
    except Exception as e:
        logger.warning(f"无法获取 PDF 页数：{e}")
        return 0


def get_docx_page_count(file_path: str) -> int:
    """
    获取 DOCX 文件页数
    
    方法优先级：
    1. 如果安装了 LibreOffice，使用命令行转换为 PDF 后获取准确页数
    2. 否则使用改进的估算法
    3. 如果都无法使用，基于文件大小估算
    """
    try:
        # 方法 1: 尝试使用 LibreOffice 转换为 PDF（最准确）
        import subprocess
        
        output_pdf = file_path + ".temp.pdf"
        cmd = [
            "soffice", "--headless", "--convert-to", "pdf",
            "--outdir", str(Path(file_path).parent),
            file_path
        ]
        
        result = subprocess.run(cmd, capture_output=True, timeout=60)
        
        if result.returncode == 0 and os.path.exists(output_pdf):
            # 使用 pdfplumber 获取准确页数
            import pdfplumber
            with pdfplumber.open(output_pdf) as pdf:
                page_count = len(pdf.pages)
            
            # 删除临时 PDF
            os.remove(output_pdf)
            logger.info(f"DOCX 页数 (LibreOffice): {page_count}页")
            return page_count
            
    except Exception as e:
        logger.debug(f"LibreOffice 转换不可用：{e}")
    
    # 方法 2: 使用改进的估算法
    try:
        from docx import Document
        
        doc = Document(file_path)
        
        # 更精确的估算：考虑更多因素
        total_lines = 0
        
        for para in doc.paragraphs:
            if para.text.strip():
                # 默认每 35 字符约一行（中文）
                char_count = len(para.text.strip())
                lines = max(1, char_count // 35)
                
                # 段落间距额外占用
                lines += 1
                total_lines += lines
        
        # 表格占用更多空间
        for table in doc.tables:
            # 每个表格至少占 3 行（标题 + 内容 + 间距）
            table_lines = 3 + len(table.rows) * 2
            total_lines += table_lines
        
        # 估算页数：每页约 30 行（考虑标题、间距等）
        estimated_pages = max(1, (total_lines + 29) // 30)
        
        logger.info(f"DOCX 页数 (估算): {estimated_pages}页 (基于{total_lines}行)")
        return estimated_pages
        
    except ImportError as e:
        logger.warning(f"python-docx 未安装，使用文件大小估算：{e}")
        # 方法 3: 基于文件大小估算（粗略）
        size_mb = get_file_size_mb(file_path)
        # 假设每页约 50KB（保守估计）
        estimated_pages = int(size_mb * 1024 / 50)
        logger.info(f"DOCX 页数 (文件大小估算): 约{estimated_pages}页 ({size_mb:.2f}MB)")
        return estimated_pages
    except Exception as e:
        logger.warning(f"无法估算 DOCX 页数：{e}")
        return 0


def needs_split(file_path: str) -> bool:
    """
    检查文件是否需要分割
    
    检查条件：
    1. 文件大小 > MAX_FILE_SIZE_MB
    2. PDF 页数 > MAX_PAGES
    3. DOCX 页数 > MAX_PAGES
    """
    ext = Path(file_path).suffix.lower()
    
    # 检查文件大小
    size_mb = get_file_size_mb(file_path)
    if size_mb > MAX_FILE_SIZE_MB:
        logger.info(f"文件大小超限：{size_mb:.2f}MB > {MAX_FILE_SIZE_MB}MB")
        return True
    
    # 检查页数（仅 PDF 和 DOCX）
    if ext == ".pdf":
        page_count = get_pdf_page_count(file_path)
        if page_count > MAX_PAGES:
            logger.info(f"PDF 页数超限：{page_count}页 > {MAX_PAGES}页")
            return True
    elif ext in [".docx", ".doc"]:
        page_count = get_docx_page_count(file_path)
        if page_count > MAX_PAGES:
            logger.info(f"DOCX 页数超限：约{page_count}页 > {MAX_PAGES}页")
            return True
    
    return False
