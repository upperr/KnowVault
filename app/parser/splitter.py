"""
文档分割模块
将超限的 PDF 或 DOCX 文件分割为多个子文件，暂存到 output 目录

使用流程：
1. 调用 split_file() 分割文件，子文件保存到 output/ 目录
2. 对每个子文件进行 MinerU 解析
3. 调用 cleanup_sub_files() 删除子文件
4. 保留原始文件和解析后的 MD 文件
"""
import os
import shutil
import logging
from pathlib import Path
from typing import List

from app.parser.config import MAX_PAGES

logger = logging.getLogger(__name__)

# 子文档暂存目录
OUTPUT_DIR = Path(__file__).parent.parent.parent / "data" / "output" / "split_files"


def get_output_dir() -> Path:
    """获取并创建 output 目录"""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    return OUTPUT_DIR


def clear_output_dir():
    """清空 output 目录（可选操作）"""
    if OUTPUT_DIR.exists():
        shutil.rmtree(OUTPUT_DIR)
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        logger.info("已清空 output 目录")


def split_pdf(file_path: str) -> List[str]:
    """
    分割 PDF 文件，返回子文件路径列表
    
    子文件保存到 data/output/split_files/ 目录
    每个子文件最多 MAX_PAGES 页
    """
    from pypdf import PdfReader, PdfWriter
    
    output_files = []
    try:
        reader = PdfReader(file_path)
        total_pages = len(reader.pages)
        
        # 计算每个子文件的页数（留一些余量）
        pages_per_file = MAX_PAGES - 10
        
        # 使用 output 目录作为子文件暂存目录
        output_dir = get_output_dir()
        base_name = Path(file_path).stem
        
        split_idx = 0
        for start_page in range(0, total_pages, pages_per_file):
            end_page = min(start_page + pages_per_file, total_pages)
            
            writer = PdfWriter()
            for i in range(start_page, end_page):
                writer.add_page(reader.pages[i])
            
            # 保存子文件到 output 目录
            output_filename = f"{base_name}_part{split_idx + 1}.pdf"
            output_path = output_dir / output_filename
            
            with open(output_path, "wb") as f:
                writer.write(f)
            
            output_files.append(str(output_path))
            logger.info(f"已分割 PDF: {output_filename} (第{start_page + 1}-{end_page}页) -> {output_path}")
            split_idx += 1
        
        logger.info(f"PDF 分割完成：共{len(output_files)}个子文件，暂存于：{output_dir}")
        return output_files
        
    except Exception as e:
        logger.error(f"PDF 分割失败：{e}")
        return []


def split_docx(file_path: str) -> List[str]:
    """
    分割 DOCX 文件，返回子文件路径列表
    
    子文件保存到 data/output/split_files/ 目录
    按段落和表格元素分割
    """
    from docx import Document
    
    output_files = []
    try:
        doc = Document(file_path)
        
        # 收集所有段落和表格
        all_elements = []
        for para in doc.paragraphs:
            all_elements.append(("para", para))
        for table in doc.tables:
            all_elements.append(("table", table))
        
        # 计算每个子文件的元素数
        elements_per_file = max(50, len(all_elements) // (MAX_PAGES // 10))
        
        # 使用 output 目录作为子文件暂存目录
        output_dir = get_output_dir()
        base_name = Path(file_path).stem
        
        split_idx = 0
        for start_idx in range(0, len(all_elements), elements_per_file):
            end_idx = min(start_idx + elements_per_file, len(all_elements))
            
            # 创建新文档
            new_doc = Document()
            
            # 复制元素
            for i in range(start_idx, end_idx):
                elem_type, elem = all_elements[i]
                if elem_type == "para":
                    new_doc.add_paragraph(elem.text)
                elif elem_type == "table":
                    # 复制表格
                    new_table = new_doc.add_table(
                        rows=len(elem.rows), 
                        cols=len(elem.columns)
                    )
                    for r, row in enumerate(elem.rows):
                        for c, cell in enumerate(row.cells):
                            new_table.cell(r, c).text = cell.text
            
            # 保存子文件到 output 目录
            output_filename = f"{base_name}_part{split_idx + 1}.docx"
            output_path = output_dir / output_filename
            
            new_doc.save(output_path)
            output_files.append(str(output_path))
            logger.info(f"已分割 DOCX: {output_filename} (元素{start_idx + 1}-{end_idx}) -> {output_path}")
            split_idx += 1
        
        logger.info(f"DOCX 分割完成：共{len(output_files)}个子文件，暂存于：{output_dir}")
        return output_files
        
    except Exception as e:
        logger.error(f"DOCX 分割失败：{e}")
        return []


def split_file(file_path: str) -> List[str]:
    """
    分割文件，返回子文件路径列表
    
    子文件保存到 data/output/split_files/ 目录
    根据文件类型自动选择分割方法
    """
    ext = Path(file_path).suffix.lower()
    
    if ext == ".pdf":
        return split_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        return split_docx(file_path)
    else:
        logger.warning(f"不支持的分割格式：{ext}")
        return []


def delete_file(file_path: str) -> bool:
    """删除文件"""
    try:
        os.remove(file_path)
        logger.info(f"已删除文件：{file_path}")
        return True
    except Exception as e:
        logger.warning(f"删除文件失败 {file_path}: {e}")
        return False


def cleanup_sub_files(sub_files: List[str]) -> int:
    """
    清理子文件（分割后的临时文件）
    
    Args:
        sub_files: 子文件路径列表
    
    Returns:
        成功删除的文件数量
    """
    deleted_count = 0
    for sub_file in sub_files:
        if delete_file(sub_file):
            deleted_count += 1
    
    if deleted_count > 0:
        logger.info(f"已清理 {deleted_count}/{len(sub_files)} 个子文件")
    
    return deleted_count
